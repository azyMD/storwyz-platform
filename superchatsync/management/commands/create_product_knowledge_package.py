import json
import os
import re
from pathlib import Path

import requests
from docx import Document

from django.conf import settings
from django.core.management.base import BaseCommand, CommandError
from django.db import connection
from django.utils import timezone


BASE_DIR = "/opt/superchat-ai-agent"
ENV_PATH = os.path.join(BASE_DIR, ".env")
PACKAGE_ROOT = "/opt/superchat-ai-agent/data/product_knowledge_packages"
DEFAULT_MODEL = "gpt-4.1-mini"


ITEM_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": [
        "title",
        "question",
        "answer",
        "rule",
        "keyword",
        "description",
        "price",
        "priority",
        "target_product_name",
        "target_product_id",
        "evidence",
        "confidence_score",
    ],
    "properties": {
        "title": {"type": "string"},
        "question": {"type": "string"},
        "answer": {"type": "string"},
        "rule": {"type": "string"},
        "keyword": {"type": "string"},
        "description": {"type": "string"},
        "price": {"type": "string"},
        "priority": {"type": "integer", "minimum": 0, "maximum": 100},
        "target_product_name": {"type": "string"},
        "target_product_id": {"type": "string"},
        "evidence": {"type": "string"},
        "confidence_score": {"type": "integer", "minimum": 0, "maximum": 100},
    },
}


COMMON_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["items"],
    "properties": {
        "items": {
            "type": "array",
            "maxItems": 60,
            "items": ITEM_SCHEMA,
        }
    },
}


CATEGORIES = [
    {
        "key": "product_facts",
        "filename": "01_product_facts.jsonl",
        "instruction": """
Extrage doar date stabile despre produsul principal:
producător, greutate, dimensiuni, material, caracteristici, beneficii concrete.
Nu extrage scripturi de vânzare.
Nu extrage cross-sell-uri.
Fiecare fapt trebuie să fie un item separat.
""",
    },
    {
        "key": "offers",
        "filename": "02_offers.jsonl",
        "instruction": """
Extrage ofertele, variantele, pachetele, prețurile, reducerile și bonusurile pentru produsul principal.
Include versiuni premium, bundle-uri, preț 1 buc, preț 2 buc, cadouri.
Fiecare ofertă/variantă trebuie să fie un item separat.
""",
    },
    {
        "key": "objection_rules",
        "filename": "03_objection_rules.jsonl",
        "instruction": """
Extrage obiecțiile clientului și răspunsul recomandat.
Exemple: nu am timp, de ce livrarea e cu plată, am îndoieli privind calitatea, e scump, vreau mai târziu.
Fiecare obiecție trebuie să fie un item separat.
În question pune obiecția clientului.
În answer pune răspunsul recomandat.
""",
    },
    {
        "key": "sales_rules",
        "filename": "04_sales_rules.jsonl",
        "instruction": """
Extrage regulile de vânzare și pașii pe care operatorul/AI trebuie să îi urmeze.
Exemple: ce întrebări trebuie puse, când se cere adresa, cum se confirmă comanda, cum se prezintă premium, când se prezintă cross-sell.
Fiecare regulă trebuie să fie un item separat.
În rule pune regula clară și aplicabilă.
""",
    },
    {
        "key": "cross_sell_rules",
        "filename": "05_cross_sell_rules.jsonl",
        "instruction": """
Extrage produsele cross-sell / upsell secundare.
Pentru fiecare produs adițional extrage: nume, descriere, preț, când se recomandă, beneficiu.
Nu le amesteca cu produsul principal.
Fiecare produs cross-sell trebuie să fie un item separat.
""",
    },
    {
        "key": "product_faq",
        "filename": "06_product_faq.jsonl",
        "instruction": """
Extrage întrebări și răspunsuri utile pentru client despre produs, ofertă, livrare, plată, utilizare, calitate.
Dacă documentul nu are întrebarea formulată explicit, creează o întrebare naturală bazată strict pe informația din document.
Fiecare FAQ trebuie să fie un item separat.
În question pune întrebarea.
În answer pune răspunsul.
""",
    },
    {
        "key": "detection_keywords",
        "filename": "07_detection_keywords.jsonl",
        "instruction": """
Extrage keyword-uri pentru detectarea produsului și a intenției clientului.
Include nume produs, denumiri alternative, termeni din reclamă, termeni populari, nume cross-sell.
Fiecare keyword trebuie să fie un item separat.
În keyword pune doar cuvântul sau expresia.
""",
    },
    {
        "key": "conversation_examples",
        "filename": "08_conversation_examples.jsonl",
        "instruction": """
Extrage exemple concrete de conversație / răspuns bun.
Include situații de tip: client nu are timp, întreabă de livrare, are îndoieli, finalizare comandă, propunere premium, cross-sell.
Fiecare exemplu trebuie să fie un item separat.
În question pune mesajul/situația clientului.
În answer pune răspunsul recomandat.
""",
    },
    {
        "key": "workflow_rules",
        "filename": "09_workflow_rules.jsonl",
        "instruction": """
Extrage reguli de workflow operațional.
Exemple: verificare nume/prenume/adresă, confirmare produse, termen livrare, notificare SMS curier, verificare culori în CRM.
Fiecare regulă trebuie să fie un item separat.
În rule pune regula operațională.
""",
    },
]


def load_env_file():
    if not os.path.exists(ENV_PATH):
        return

    with open(ENV_PATH, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()

            if not line or line.startswith("#") or "=" not in line:
                continue

            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def dictfetchall(cur):
    cols = [c[0] for c in cur.description]
    return [dict(zip(cols, row)) for row in cur.fetchall()]


def get_import(import_id):
    with connection.cursor() as cur:
        cur.execute("""
            SELECT *
            FROM product_knowledge_imports
            WHERE import_id = %s
        """, [str(import_id)])
        rows = dictfetchall(cur)

    return rows[0] if rows else None


def update_import(import_id, **fields):
    fields["updated_at"] = timezone.now()

    set_sql = ", ".join([f"{k} = %s" for k in fields.keys()])
    params = list(fields.values()) + [str(import_id)]

    with connection.cursor() as cur:
        cur.execute(f"""
            UPDATE product_knowledge_imports
            SET {set_sql}
            WHERE import_id = %s
        """, params)


def table_exists(table_name):
    with connection.cursor() as cur:
        cur.execute("""
            SELECT EXISTS (
                SELECT 1 FROM information_schema.tables WHERE table_name = %s
            )
        """, [table_name])
        return cur.fetchone()[0]


def fetch_table_records(table_name, product_id=None, limit=80):
    if not table_exists(table_name):
        return []

    with connection.cursor() as cur:
        cur.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_name = %s
        """, [table_name])
        columns = {r[0] for r in cur.fetchall()}

    where = ""
    params = []

    if product_id and "product_id" in columns:
        where = "WHERE product_id = %s"
        params.append(product_id)

    sql = f"SELECT * FROM {table_name} {where} LIMIT %s"
    params.append(limit)

    with connection.cursor() as cur:
        cur.execute(sql, params)
        rows = dictfetchall(cur)

    clean_rows = []
    for row in rows:
        clean = {}
        for k, v in row.items():
            if hasattr(v, "isoformat"):
                clean[k] = v.isoformat()
            else:
                clean[k] = v
        clean_rows.append(clean)

    return clean_rows


def get_existing_context(product_id):
    return {
        "product": fetch_table_records("products", product_id, 1),
        "offers": fetch_table_records("offers", product_id, 50),
        "product_faq": fetch_table_records("product_faq", product_id, 80),
        "product_sales_rules": fetch_table_records("product_sales_rules", product_id, 80),
        "objection_rules": fetch_table_records("objection_rules", product_id, 80),
        "cross_sell_rules": fetch_table_records("cross_sell_rules", product_id, 80),
        "product_detection_rules": fetch_table_records("product_detection_rules", product_id, 80),
    }


def resolve_file_path(source_file):
    value = str(source_file or "")

    if not value:
        raise RuntimeError("source_file este gol.")

    path = Path(value)

    if path.is_absolute():
        return path

    return Path(settings.MEDIA_ROOT) / value


def extract_docx_text(path):
    doc = Document(str(path))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    cells.append(value)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def safe_slug(value):
    value = str(value or "").lower()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    value = value.strip("_")
    return value or "knowledge"


def extract_response_text(data):
    if data.get("output_text"):
        return data["output_text"]

    texts = []

    for item in data.get("output", []):
        for c in item.get("content", []):
            if isinstance(c, dict):
                if c.get("text"):
                    texts.append(c["text"])
                elif c.get("output_text"):
                    texts.append(c["output_text"])

    return "\n".join(texts).strip()


def call_openai_category(api_key, model, product_id, import_title, category, text, existing_context):
    system = f"""
Ești un AI specializat în normalizarea documentelor comerciale în knowledge package pentru AI sales agent.

Documentul poate fi:
- script de apel
- FAQ
- landing page
- catalog
- ofertă
- combinație dintre ele

Sarcina ta este strict pentru categoria: {category["key"]}

Instrucțiuni categorie:
{category["instruction"]}

Reguli generale:
1. Nu inventa informații.
2. Nu face rezumat general.
3. Extrage itemi atomici separați.
4. Dacă informația nu există pentru categoria dată, returnează array gol.
5. Evidence trebuie să indice fragmentul din document care susține itemul.
6. Scrie în română.
7. Răspunde strict JSON conform schemei.
"""

    payload = {
        "product_id": product_id,
        "import_title": import_title,
        "category": category["key"],
        "document_text": text[:90000],
        "existing_product_context": existing_context,
    }

    body = {
        "model": model,
        "instructions": system,
        "input": json.dumps(payload, ensure_ascii=False),
        "text": {
            "format": {
                "type": "json_schema",
                "name": f"knowledge_package_{category['key']}",
                "strict": True,
                "schema": COMMON_SCHEMA,
            }
        },
    }

    response = requests.post(
        "https://api.openai.com/v1/responses",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=180,
    )

    if response.status_code >= 400:
        raise RuntimeError(f"OpenAI error {response.status_code}: {response.text[:1500]}")

    data = response.json()
    text_response = extract_response_text(data)

    if not text_response:
        return [], data.get("id")

    parsed = json.loads(text_response)
    return parsed.get("items", []), data.get("id")


def write_jsonl(path, items):
    with open(path, "w", encoding="utf-8") as f:
        for item in items:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")


class Command(BaseCommand):
    help = "Create structured AI knowledge package from uploaded product DOCX."

    def add_arguments(self, parser):
        parser.add_argument("--import-id", required=True)

    def handle(self, *args, **options):
        load_env_file()

        api_key = os.getenv("OPENAI_API_KEY")
        model = os.getenv("OPENAI_MODEL", DEFAULT_MODEL)

        if not api_key:
            raise CommandError("OPENAI_API_KEY lipsește din .env")

        import_id = options["import_id"]
        item = get_import(import_id)

        if not item:
            raise CommandError(f"Importul nu există: {import_id}")

        product_id = item.get("product_id")

        if not product_id:
            raise CommandError("Importul nu are product_id.")

        update_import(
            import_id,
            knowledge_package_status="processing",
            knowledge_package_error=None,
        )

        try:
            source_file = resolve_file_path(item.get("source_file"))

            if not source_file.exists():
                raise RuntimeError(f"Fișierul nu există: {source_file}")

            if source_file.suffix.lower() != ".docx":
                raise RuntimeError("Fișierul trebuie să fie .docx.")

            raw_text = extract_docx_text(source_file)

            if len(raw_text) < 100:
                raise RuntimeError("Text extras prea scurt.")

            package_name = f"{safe_slug(product_id)}_{safe_slug(item.get('title'))}_{str(import_id)[:8]}"
            package_dir = Path(PACKAGE_ROOT) / package_name
            package_dir.mkdir(parents=True, exist_ok=True)

            with open(package_dir / "raw_extracted_text.txt", "w", encoding="utf-8") as f:
                f.write(raw_text)

            existing_context = get_existing_context(product_id)

            manifest = {
                "import_id": str(import_id),
                "product_id": product_id,
                "title": item.get("title"),
                "source_file": str(source_file),
                "package_dir": str(package_dir),
                "model": model,
                "created_at": timezone.now().isoformat(),
                "categories": {},
            }

            total_items = 0

            for category in CATEGORIES:
                self.stdout.write(f"Processing category: {category['key']}")

                items, response_id = call_openai_category(
                    api_key=api_key,
                    model=model,
                    product_id=product_id,
                    import_title=item.get("title"),
                    category=category,
                    text=raw_text,
                    existing_context=existing_context,
                )

                file_path = package_dir / category["filename"]
                write_jsonl(file_path, items)

                manifest["categories"][category["key"]] = {
                    "filename": category["filename"],
                    "count": len(items),
                    "openai_response_id": response_id,
                }

                total_items += len(items)

                self.stdout.write(self.style.SUCCESS(
                    f"{category['key']}: {len(items)} items"
                ))

            with open(package_dir / "00_manifest.json", "w", encoding="utf-8") as f:
                json.dump(manifest, f, ensure_ascii=False, indent=2)

            with open(package_dir / "source_summary.md", "w", encoding="utf-8") as f:
                f.write(f"# Knowledge Package\n\n")
                f.write(f"Product ID: {product_id}\n\n")
                f.write(f"Import ID: {import_id}\n\n")
                f.write(f"Title: {item.get('title')}\n\n")
                f.write(f"Total extracted items: {total_items}\n\n")
                f.write("## Category counts\n\n")
                for key, data in manifest["categories"].items():
                    f.write(f"- {key}: {data['count']}\n")

            update_import(
                import_id,
                knowledge_package_status="created",
                knowledge_package_dir=str(package_dir),
                knowledge_package_error=None,
                knowledge_package_created_at=timezone.now(),
                extracted_text=raw_text,
                extracted_char_count=len(raw_text),
            )

            self.stdout.write(self.style.SUCCESS(
                f"Knowledge package created: {package_dir} | total_items={total_items}"
            ))

        except Exception as e:
            update_import(
                import_id,
                knowledge_package_status="error",
                knowledge_package_error=str(e),
            )
            self.stdout.write(self.style.ERROR(str(e)))
            raise
