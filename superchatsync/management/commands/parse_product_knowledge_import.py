import json
import os
from pathlib import Path

import requests
from docx import Document

from django.conf import settings
from django.core.management.base import BaseCommand, CommandError
from django.db import connection
from django.utils import timezone


BASE_DIR = "/opt/superchat-ai-agent"
ENV_PATH = os.path.join(BASE_DIR, ".env")
DEFAULT_MODEL = "gpt-4.1-mini"
CHUNK_SIZE = 14000


SUGGESTION_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["suggestions"],
    "properties": {
        "suggestions": {
            "type": "array",
            "maxItems": 50,
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "suggestion_type",
                    "title",
                    "suggested_question",
                    "suggested_answer",
                    "suggested_rule",
                    "suggested_keyword",
                    "reason",
                    "evidence",
                    "confidence_score",
                ],
                "properties": {
                    "suggestion_type": {
                        "type": "string",
                        "enum": [
                            "product_faq",
                            "objection_rule",
                            "sales_rule",
                            "detection_keyword",
                            "conversation_example",
                            "workflow_fix",
                        ],
                    },
                    "title": {"type": "string"},
                    "suggested_question": {"type": "string"},
                    "suggested_answer": {"type": "string"},
                    "suggested_rule": {"type": "string"},
                    "suggested_keyword": {"type": "string"},
                    "reason": {"type": "string"},
                    "evidence": {"type": "string"},
                    "confidence_score": {
                        "type": "integer",
                        "minimum": 0,
                        "maximum": 100,
                    },
                },
            },
        }
    },
}


SYSTEM_INSTRUCTIONS = """
Ești un specialist în transformarea documentelor comerciale în product knowledge base pentru AI sales agent.

Primești fragmente dintr-un document .docx despre un produs.
Scopul tău NU este să faci rezumat.
Scopul tău este să extragi cât mai multe elemente atomice, clare și aplicabile în product feed.

Trebuie să generezi multe sugestii utile, dacă documentul conține informație suficientă.

Tipuri:
1. product_faq
   Folosește pentru întrebări și răspunsuri concrete.
   Exemplu: clientul întreabă compatibilitate, livrare, plată, material, utilizare, garanție.

2. objection_rule
   Folosește pentru obiecții de client: scump, nu am încredere, nu știu dacă se potrivește, vreau mai târziu, am mai văzut, etc.

3. sales_rule
   Folosește pentru reguli de vânzare:
   - ce să întrebe operatorul
   - când să ceară adresa
   - cum să confirme comanda
   - ce să nu promită
   - cum să răspundă la interes clar

4. detection_keyword
   Folosește pentru denumiri alternative ale produsului, keyword-uri din reclame, nume populare, greșeli uzuale de scriere.

5. conversation_example
   Folosește pentru exemple de răspuns bun către client.

6. workflow_fix
   Folosește doar dacă documentul indică o modificare clară de workflow.

Reguli foarte importante:
- Nu returna doar 1 sugestie generală dacă documentul conține mai multe informații.
- Nu face rezumat. Creează itemi separați.
- Nu inventa specificații.
- Dacă nu ești sigur, transformă informația în sales_rule de clarificare.
- Dacă documentul conține scripturi de conversație, transformă fiecare situație utilă în FAQ / objection_rule / sales_rule / conversation_example.
- Dacă documentul conține oferte/prețuri, creează FAQ sau sales_rule despre cum se comunică oferta.
- Dacă documentul conține mesaje recomandate, creează conversation_example.
- Evidence trebuie să fie un scurt citat/parafrază din document.
- Răspunde strict JSON conform schemei.
"""


def load_env_file():
    if not os.path.exists(ENV_PATH):
        return

    with open(ENV_PATH, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()

            if not line or line.startswith("#") or "=" not in line:
                continue

            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def dictfetchall(cur):
    cols = [c[0] for c in cur.description]
    return [dict(zip(cols, row)) for row in cur.fetchall()]


def table_exists(table_name):
    with connection.cursor() as cur:
        cur.execute("""
            SELECT EXISTS (
                SELECT 1
                FROM information_schema.tables
                WHERE table_name = %s
            )
        """, [table_name])
        return cur.fetchone()[0]


def fetch_table_records(table_name, product_id=None, limit=80):
    if not table_exists(table_name):
        return []

    with connection.cursor() as cur:
        cur.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_name = %s
        """, [table_name])
        columns = {row[0] for row in cur.fetchall()}

    where = ""
    params = []

    if product_id and "product_id" in columns:
        where = "WHERE product_id = %s"
        params.append(product_id)

    sql = f"SELECT * FROM {table_name} {where} LIMIT %s"
    params.append(limit)

    with connection.cursor() as cur:
        cur.execute(sql, params)
        rows = dictfetchall(cur)

    clean_rows = []

    for row in rows:
        clean = {}
        for k, v in row.items():
            if hasattr(v, "isoformat"):
                clean[k] = v.isoformat()
            else:
                clean[k] = v
        clean_rows.append(clean)

    return clean_rows


def get_import(import_id):
    with connection.cursor() as cur:
        cur.execute("""
            SELECT
                import_id,
                product_id,
                title,
                source_file,
                original_filename,
                status,
                notes
            FROM product_knowledge_imports
            WHERE import_id = %s
        """, [str(import_id)])
        rows = dictfetchall(cur)

    return rows[0] if rows else None


def update_import(import_id, **fields):
    if not fields:
        return

    fields["updated_at"] = timezone.now()

    set_sql = ", ".join([f"{key} = %s" for key in fields.keys()])
    params = list(fields.values()) + [str(import_id)]

    with connection.cursor() as cur:
        cur.execute(f"""
            UPDATE product_knowledge_imports
            SET {set_sql}
            WHERE import_id = %s
        """, params)


def extract_docx_text(path):
    doc = Document(str(path))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    cells.append(value)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def split_text(text, size=CHUNK_SIZE):
    text = text.strip()

    if len(text) <= size:
        return [text]

    chunks = []
    current = []

    current_len = 0

    for paragraph in text.split("\n"):
        p = paragraph.strip()
        if not p:
            continue

        if current_len + len(p) > size and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0

        current.append(p)
        current_len += len(p) + 1

    if current:
        chunks.append("\n".join(current))

    return chunks


def resolve_file_path(source_file):
    value = str(source_file or "")

    if not value:
        raise RuntimeError("source_file este gol.")

    path = Path(value)

    if path.is_absolute():
        return path

    media_root = Path(getattr(settings, "MEDIA_ROOT", "/opt/superchat-ai-agent/media"))
    return media_root / value


def get_product_context(product_id):
    return {
        "product": fetch_table_records("products", product_id, 1),
        "offers": fetch_table_records("offers", product_id, 50),
        "product_faq": fetch_table_records("product_faq", product_id, 80),
        "product_sales_rules": fetch_table_records("product_sales_rules", product_id, 80),
        "objection_rules": fetch_table_records("objection_rules", product_id, 80),
        "product_detection_rules": fetch_table_records("product_detection_rules", product_id, 80),
        "cross_sell_rules": fetch_table_records("cross_sell_rules", product_id, 80),
    }


def extract_response_text(data):
    if data.get("output_text"):
        return data["output_text"]

    texts = []

    for item in data.get("output", []):
        for c in item.get("content", []):
            if isinstance(c, dict):
                if c.get("text"):
                    texts.append(c["text"])
                elif c.get("output_text"):
                    texts.append(c["output_text"])

    return "\n".join(texts).strip()


def call_openai(api_key, model, payload):
    body = {
        "model": model,
        "instructions": SYSTEM_INSTRUCTIONS,
        "input": json.dumps(payload, ensure_ascii=False),
        "text": {
            "format": {
                "type": "json_schema",
                "name": "product_knowledge_import_suggestions",
                "strict": True,
                "schema": SUGGESTION_SCHEMA,
            }
        },
    }

    response = requests.post(
        "https://api.openai.com/v1/responses",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=180,
    )

    if response.status_code >= 400:
        raise RuntimeError(f"OpenAI error {response.status_code}: {response.text[:1500]}")

    data = response.json()
    text = extract_response_text(data)

    if not text:
        raise RuntimeError("OpenAI response nu conține text.")

    return json.loads(text), data


def suggestion_exists(product_id, suggestion_type, title, question, rule, keyword):
    with connection.cursor() as cur:
        cur.execute("""
            SELECT 1
            FROM product_feed_suggestions
            WHERE product_id = %s
              AND suggestion_type = %s
              AND status IN ('pending_review', 'approved', 'applied')
              AND (
                    COALESCE(title, '') = COALESCE(%s, '')
                 OR COALESCE(suggested_question, '') = COALESCE(%s, '')
                 OR COALESCE(suggested_rule, '') = COALESCE(%s, '')
                 OR COALESCE(suggested_keyword, '') = COALESCE(%s, '')
              )
            LIMIT 1
        """, [
            product_id,
            suggestion_type,
            title,
            question,
            rule,
            keyword,
        ])
        return cur.fetchone() is not None


def delete_existing_import_suggestions(import_id):
    with connection.cursor() as cur:
        cur.execute("""
            DELETE FROM product_feed_suggestions
            WHERE created_by = 'ai_knowledge_import'
              AND raw_payload->>'import_id' = %s
              AND status IN ('pending_review', 'rejected')
        """, [str(import_id)])


def save_suggestions(import_id, product_id, suggestions, raw_response, chunk_index):
    inserted = 0
    skipped = 0

    with connection.cursor() as cur:
        for item in suggestions:
            confidence = item.get("confidence_score") or 0

            if confidence < 45:
                skipped += 1
                continue

            suggestion_type = item.get("suggestion_type")
            title = item.get("title") or ""
            question = item.get("suggested_question") or ""
            answer = item.get("suggested_answer") or ""
            rule = item.get("suggested_rule") or ""
            keyword = item.get("suggested_keyword") or ""

            if suggestion_exists(product_id, suggestion_type, title, question, rule, keyword):
                skipped += 1
                continue

            cur.execute("""
                INSERT INTO product_feed_suggestions (
                    conversation_id,
                    analysis_id,
                    product_id,
                    suggestion_type,
                    title,
                    suggested_question,
                    suggested_answer,
                    suggested_rule,
                    suggested_keyword,
                    reason,
                    evidence,
                    confidence_score,
                    status,
                    created_by,
                    raw_payload,
                    created_at,
                    updated_at
                )
                VALUES (
                    NULL,
                    NULL,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    %s,
                    'pending_review',
                    'ai_knowledge_import',
                    %s::jsonb,
                    NOW(),
                    NOW()
                )
            """, [
                product_id,
                suggestion_type,
                title,
                question,
                answer,
                rule,
                keyword,
                item.get("reason"),
                item.get("evidence"),
                confidence,
                json.dumps({
                    "source": "product_knowledge_import",
                    "import_id": str(import_id),
                    "chunk_index": chunk_index,
                    "openai_response_id": raw_response.get("id"),
                    "suggestion": item,
                }, ensure_ascii=False),
            ])

            inserted += 1

    return inserted, skipped


class Command(BaseCommand):
    help = "Parse uploaded .docx product knowledge document and create many Product Feed Suggestions."

    def add_arguments(self, parser):
        parser.add_argument("--import-id", required=True)
        parser.add_argument("--force", action="store_true")

    def handle(self, *args, **options):
        load_env_file()

        import_id = options["import_id"]
        force = options.get("force")
        api_key = os.getenv("OPENAI_API_KEY")
        model = os.getenv("OPENAI_MODEL", DEFAULT_MODEL)

        if not api_key:
            raise CommandError("OPENAI_API_KEY lipsește din /opt/superchat-ai-agent/.env")

        item = get_import(import_id)

        if not item:
            raise CommandError(f"ProductKnowledgeImport nu există: {import_id}")

        product_id = item.get("product_id")

        if not product_id:
            update_import(
                import_id,
                status="error",
                error="Importul nu are product_id. Selectează produsul în Admin.",
                processed_at=timezone.now(),
            )
            raise CommandError("Importul nu are product_id. Selectează produsul în Admin.")

        if force:
            delete_existing_import_suggestions(import_id)

        update_import(import_id, status="processing", error=None)

        try:
            file_path = resolve_file_path(item.get("source_file"))

            if not file_path.exists():
                raise RuntimeError(f"Fișierul nu există: {file_path}")

            if file_path.suffix.lower() != ".docx":
                raise RuntimeError("Fișierul trebuie să fie .docx. Exportă din Google Docs ca Microsoft Word (.docx).")

            text = extract_docx_text(file_path)

            if len(text) < 100:
                raise RuntimeError("Text extras prea scurt. Verifică dacă documentul conține text real.")

            chunks = split_text(text)
            product_context = get_product_context(product_id)

            total_inserted = 0
            total_skipped = 0

            for idx, chunk in enumerate(chunks, start=1):
                payload = {
                    "product_id": product_id,
                    "import_title": item.get("title"),
                    "chunk_index": idx,
                    "chunks_total": len(chunks),
                    "document_chunk_text": chunk,
                    "existing_product_feed": product_context,
                    "task": {
                        "extract_many_atomic_items": True,
                        "create_product_feed_suggestions": True,
                        "do_not_modify_product_feed_directly": True,
                        "minimum_expected_items_if_content_exists": 8,
                    }
                }

                result, raw_response = call_openai(api_key, model, payload)

                inserted, skipped = save_suggestions(
                    import_id=import_id,
                    product_id=product_id,
                    suggestions=result.get("suggestions", []),
                    raw_response=raw_response,
                    chunk_index=idx,
                )

                total_inserted += inserted
                total_skipped += skipped

                self.stdout.write(self.style.SUCCESS(
                    f"Chunk {idx}/{len(chunks)}: inserted={inserted}, skipped={skipped}"
                ))

            update_import(
                import_id,
                status="suggestions_created",
                extracted_text=text,
                extracted_char_count=len(text),
                suggestions_created_count=total_inserted,
                error=None,
                processed_at=timezone.now(),
            )

            self.stdout.write(self.style.SUCCESS(
                f"Import parsed. product_id={product_id}, chunks={len(chunks)}, inserted={total_inserted}, skipped={total_skipped}"
            ))

        except Exception as e:
            update_import(
                import_id,
                status="error",
                error=str(e),
                processed_at=timezone.now(),
            )
            self.stdout.write(self.style.ERROR(str(e)))
            raise
